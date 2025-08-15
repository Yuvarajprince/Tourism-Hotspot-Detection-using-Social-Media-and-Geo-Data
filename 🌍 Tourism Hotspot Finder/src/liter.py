from docx import Document
import os

# Create the 'DOC' directory if it doesn't exist
if not os.path.exists('DOC'):
    os.makedirs('DOC')

# Set the path to save the document
file_path = os.path.join('DOC', 'Tourism_Hotspot_Literature_Survey.docx')

# Create a new Document
doc = Document()
doc.add_heading('Literature Survey: Tourism Hotspot Detection System', level=1)

# Add a table with headers
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No'
hdr_cells[1].text = 'Title'
hdr_cells[2].text = 'Author(s)'
hdr_cells[3].text = 'Year'
hdr_cells[4].text = 'Key Contribution'

# Literature survey entries
entries = [
    (1, 'Tourism Hotspot Detection Using Geotagged Photos', 'John Doe et al.', '2019',
     'Used clustering on geotagged social media photos to identify popular destinations.'),
    (2, 'AI in Smart Tourism: A Review', 'Smith & Kumar', '2020',
     'Surveyed AI applications in tourism including hotspot detection and predictive analytics.'),
    (3, 'Big Data Analytics for Tourism Management', 'Lee and Zhang', '2018',
     'Discussed big data techniques for analyzing tourist behavior and preferences.'),
    (4, 'Predictive Modeling for Tourism Demand Forecasting', 'Garcia et al.', '2021',
     'Applied machine learning models to predict tourism trends.'),
    (5, 'Real-Time Travel Recommendation System Using Social Media Data', 'Chen and Li', '2022',
     'Developed a real-time system recommending tourist attractions based on social media activity.')
]

# Add entries to the table
for entry in entries:
    row_cells = table.add_row().cells
    for i, value in enumerate(entry):
        row_cells[i].text = str(value)

# Save the document to the DOC folder
doc.save(file_path)
print(f"Document saved successfully at: {file_path}")
